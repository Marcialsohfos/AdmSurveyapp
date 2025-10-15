import pandas as pd
import os
import json

class DataConverter:
    def __init__(self):
        self.output_dir = 'converted'
        os.makedirs(self.output_dir, exist_ok=True)
    
    def convert_data(self, data, output_format):
        """Convertit les données dans le format demandé"""
        print(f"🔄 Conversion demandée: format {output_format}")
        print(f"📊 Données reçues: {type(data)}")
        
        filename = f"exported_data.{output_format}"
        filepath = os.path.join(self.output_dir, filename)
        
        print(f"💾 Fichier de sortie: {filepath}")
        
        # Vérifier que le dossier existe
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Conversion selon le format
        try:
            if output_format == 'docx':
                result = self._to_docx(data, filepath)
            elif output_format == 'txt':
                result = self._to_txt(data, filepath)
            elif output_format == 'csv':
                result = self._to_csv(data, filepath)
            elif output_format == 'xlsx':
                result = self._to_excel(data, filepath)
            elif output_format == 'json':
                result = self._to_json(data, filepath)
            elif output_format == 'dta':
                result = self._to_stata(data, filepath)
            elif output_format == 'sav':
                result = self._to_spss(data, filepath)
            elif output_format == 'shp':
                result = self._to_shapefile(data, filepath)
            elif output_format == 'xml':
                result = self._to_xml(data, filepath)
            elif output_format == 'html':
                result = self._to_html(data, filepath)
            else:
                # Fallback vers CSV
                result = self._to_csv(data, filepath)
            
            # Vérifier que le fichier a été créé
            if os.path.exists(filepath):
                print(f"✅ Fichier créé avec succès: {filepath}")
                return filename
            else:
                print(f"❌ Fichier non créé: {filepath}")
                # Créer un fichier de secours
                fallback_file = self._create_fallback_file(data, output_format)
                return fallback_file
                
        except Exception as e:
            print(f"❌ Erreur lors de la conversion: {str(e)}")
            import traceback
            traceback.print_exc()
            # Créer un fichier de secours avec les données brutes
            fallback_file = self._create_fallback_file(data, output_format)
            return fallback_file

    def _create_fallback_file(self, data, original_format):
        """Crée un fichier de secours avec les données disponibles"""
        try:
            filename = f"fallback_data.txt"
            filepath = os.path.join(self.output_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=== DONNÉES EXTRACTES (FALLBACK) ===\n\n")
                f.write(f"Format original demandé: {original_format}\n")
                f.write(f"Type de données: {data.get('type', 'inconnu')}\n")
                f.write(f"OCR réussi: {data.get('ocr_success', 'inconnu')}\n\n")
                
                if data.get('raw_text'):
                    f.write("TEXTE EXTRAIT:\n")
                    f.write("=" * 50 + "\n")
                    f.write(data['raw_text'])
                    f.write("\n" + "=" * 50 + "\n")
                else:
                    f.write("Aucun texte extrait disponible.\n")
            
            print(f"🔄 Fichier de secours créé: {filename}")
            return filename
        except Exception as e:
            print(f"❌ Erreur création fichier secours: {e}")
            return "error.txt"

    def _to_docx(self, data, filepath):
        """Conversion en Word (.docx)"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Titre
            doc.add_heading('Données Extraites - SourceAdApp', 0)
            
            # Métadonnées
            if isinstance(data, dict):
                doc.add_heading('Métadonnées', level=1)
                for key, value in data.items():
                    if key != 'raw_text' and not isinstance(value, (list, dict)):
                        doc.add_paragraph(f"{key}: {value}")
            
            # Contenu texte
            if data.get('raw_text'):
                doc.add_heading('Contenu Extrait', level=1)
                # Limiter la longueur du texte pour Word
                raw_text = data['raw_text']
                if len(raw_text) > 10000:
                    raw_text = raw_text[:10000] + "... [texte tronqué]"
                doc.add_paragraph(raw_text)
            
            # Sections structurées
            if data.get('sections'):
                doc.add_heading('Sections Structurées', level=1)
                for section in data['sections']:
                    doc.add_heading(section.get('title', 'Sans titre'), level=2)
                    for content in section.get('content', []):
                        doc.add_paragraph(str(content))
            
            doc.save(filepath)
            return True
            
        except ImportError:
            # Fallback vers texte si python-docx n'est pas disponible
            print("⚠️ python-docx non disponible, fallback vers txt")
            fallback_path = filepath.replace('.docx', '.txt')
            self._to_txt(data, fallback_path)
            return False
        except Exception as e:
            print(f"❌ Erreur conversion DOCX: {e}")
            return False

    def _to_txt(self, data, filepath):
        """Conversion en texte brut (.txt)"""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=== DONNÉES EXTRAITES - SourceAdApp ===\n\n")
                
                if isinstance(data, dict):
                    # Métadonnées
                    f.write("MÉTADONNÉES:\n")
                    f.write(f"Type détecté: {data.get('type', 'Inconnu')}\n")
                    f.write(f"Type de document: {data.get('detected_type', 'Inconnu')}\n")
                    f.write(f"OCR réussi: {data.get('ocr_success', 'Inconnu')}\n")
                    f.write(f"Longueur texte: {data.get('text_length', 0)} caractères\n\n")
                    
                    # Contenu texte brut
                    if data.get('raw_text'):
                        f.write("CONTENU TEXTE:\n")
                        f.write("=" * 50 + "\n")
                        f.write(data['raw_text'])
                        f.write("\n" + "=" * 50 + "\n\n")
                    
                    # Sections structurées
                    if data.get('sections'):
                        f.write("SECTIONS STRUCTURÉES:\n")
                        for i, section in enumerate(data['sections'], 1):
                            f.write(f"\n--- Section {i}: {section.get('title', 'Sans titre')} ---\n")
                            for content in section.get('content', []):
                                f.write(f"{content}\n")
                    
                    # Tableaux détectés
                    if data.get('tables'):
                        f.write("\nTABLEAUX DÉTECTÉS:\n")
                        for i, table in enumerate(data['tables'], 1):
                            f.write(f"\nTableau {i}:\n")
                            if table.get('headers'):
                                f.write(" | ".join(table['headers']) + "\n")
                                f.write("-" * 50 + "\n")
                            for row in table.get('rows', []):
                                f.write(" | ".join(str(cell) for cell in row) + "\n")
                else:
                    f.write(str(data))
            return True
        except Exception as e:
            print(f"❌ Erreur conversion TXT: {e}")
            return False
    
    def _to_csv(self, data, filepath):
        """Conversion en CSV"""
        try:
            df = self._universal_to_dataframe(data)
            df.to_csv(filepath, index=False, encoding='utf-8')
            return True
        except Exception as e:
            print(f"❌ Erreur conversion CSV: {e}")
            return False
    
    def _to_excel(self, data, filepath):
        """Conversion en Excel avec plusieurs onglets"""
        try:
            df = self._universal_to_dataframe(data)
            
            # Créer un fichier Excel avec plusieurs onglets pour les données structurées
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Donnees_Principales', index=False)
                
                # Ajouter des onglets supplémentaires pour les données structurées
                if isinstance(data, dict):
                    # Onglet pour les métadonnées
                    if 'metadata' in data:
                        metadata_data = []
                        for key, values in data['metadata'].items():
                            if isinstance(values, list):
                                for value in values:
                                    metadata_data.append({'Type': key, 'Valeur': str(value)})
                            else:
                                metadata_data.append({'Type': key, 'Valeur': str(values)})
                        
                        if metadata_data:
                            metadata_df = pd.DataFrame(metadata_data)
                            metadata_df.to_excel(writer, sheet_name='Metadonnees', index=False)
                    
                    # Onglet spécifique pour les données RH
                    if data.get('type') == 'rh_laboratoire' and data.get('personnel_par_grade'):
                        rh_data = []
                        for item in data['personnel_par_grade']:
                            rh_data.append({
                                'Grade': item.get('grade', ''),
                                'Effectif': item.get('effectif', '')
                            })
                        
                        if rh_data:
                            rh_df = pd.DataFrame(rh_data)
                            rh_df.to_excel(writer, sheet_name='Personnel_RH', index=False)
                    
                    # Onglet pour les statistiques
                    if data.get('statistiques'):
                        stats_data = []
                        for key, value in data['statistiques'].items():
                            stats_data.append({'Statistique': key, 'Valeur': str(value)})
                        
                        if stats_data:
                            stats_df = pd.DataFrame(stats_data)
                            stats_df.to_excel(writer, sheet_name='Statistiques', index=False)
            
            return True
        except Exception as e:
            print(f"❌ Erreur conversion Excel: {e}")
            return False
    
    def _to_json(self, data, filepath):
        """Conversion en JSON"""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"❌ Erreur conversion JSON: {e}")
            return False
    
    def _to_stata(self, data, filepath):
        """Conversion en Stata (.dta)"""
        try:
            import pyreadstat
            df = self._universal_to_dataframe(data)
            pyreadstat.write_dta(df, filepath)
            return True
        except ImportError:
            # Fallback vers CSV si pyreadstat n'est pas disponible
            print("⚠️ pyreadstat non disponible, fallback vers CSV")
            fallback_path = filepath.replace('.dta', '.csv')
            self._to_csv(data, fallback_path)
            return False
        except Exception as e:
            print(f"❌ Erreur conversion Stata: {e}")
            return False
    
    def _to_spss(self, data, filepath):
        """Conversion en SPSS (.sav)"""
        try:
            import pyreadstat
            df = self._universal_to_dataframe(data)
            pyreadstat.write_sav(df, filepath)
            return True
        except ImportError:
            # Fallback vers CSV si pyreadstat n'est pas disponible
            print("⚠️ pyreadstat non disponible, fallback vers CSV")
            fallback_path = filepath.replace('.sav', '.csv')
            self._to_csv(data, fallback_path)
            return False
        except Exception as e:
            print(f"❌ Erreur conversion SPSS: {e}")
            return False
    
    def _to_shapefile(self, data, filepath):
        """Conversion en Shapefile (.shp)"""
        try:
            import geopandas as gpd
            from shapely.geometry import Point
            
            df = self._universal_to_dataframe(data)
            
            # Créer des géométries simulées basées sur les données
            geometries = []
            for i in range(len(df)):
                # Générer des points aléatoires pour la démonstration
                lat = 4.0 + (i * 0.01) % 1.0  # Simulation latitude
                lon = 11.0 + (i * 0.01) % 1.0  # Simulation longitude
                geometries.append(Point(lon, lat))
            
            # Créer le GeoDataFrame
            gdf = gpd.GeoDataFrame(df, geometry=geometries, crs="EPSG:4326")
            
            # Sauvegarder le shapefile
            gdf.to_file(filepath)
            return True
            
        except ImportError:
            # Fallback vers CSV si geopandas n'est pas disponible
            print("⚠️ geopandas non disponible, fallback vers CSV")
            fallback_path = filepath.replace('.shp', '.csv')
            self._to_csv(data, fallback_path)
            return False
        except Exception as e:
            print(f"❌ Erreur conversion Shapefile: {e}")
            return False
    
    def _to_xml(self, data, filepath):
        """Conversion en XML"""
        try:
            df = self._universal_to_dataframe(data)
            
            xml_content = '<?xml version="1.0" encoding="UTF-8"?>\n'
            xml_content += '<data>\n'
            
            for _, row in df.iterrows():
                xml_content += '  <record>\n'
                for col_name, value in row.items():
                    if pd.notna(value):
                        xml_content += f'    <{col_name}>{value}</{col_name}>\n'
                xml_content += '  </record>\n'
            
            xml_content += '</data>'
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(xml_content)
            return True
        except Exception as e:
            print(f"❌ Erreur conversion XML: {e}")
            return False
    
    def _to_html(self, data, filepath):
        """Conversion en HTML"""
        try:
            df = self._universal_to_dataframe(data)
            
            html_content = """
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>Données Extraites</title>
                <style>
                    table { border-collapse: collapse; width: 100%; }
                    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                    th { background-color: #f2f2f2; }
                    tr:nth-child(even) { background-color: #f9f9f9; }
                </style>
            </head>
            <body>
                <h1>Données Extraites par SourceAdApp</h1>
            """
            
            html_content += df.to_html(classes='table table-striped', index=False, escape=False)
            html_content += "\n</body>\n</html>"
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return True
        except Exception as e:
            print(f"❌ Erreur conversion HTML: {e}")
            return False
    
    def _universal_to_dataframe(self, data):
        """Convertit n'importe quelle structure de données en DataFrame"""
        if isinstance(data, dict) and data.get('type') == 'universal':
            return self._universal_structure_to_df(data)
        else:
            # Utiliser la méthode générique qui existe déjà
            return self._generic_to_df(data)
    
    def _universal_structure_to_df(self, data):
        """Version améliorée avec gestion des échecs OCR"""
        print(f"🛡️  Traitement des données - OCR success: {data.get('ocr_success', True)}")
        
        rows = []
        
        # Cas où l'OCR a échoué
        if not data.get('ocr_success', True):
            rows.append({
                'Category': 'System',
                'Type': 'Error',
                'Content': 'OCR a échoué - Aucun texte extrait de l\'image'
            })
            return pd.DataFrame(rows)
        
        # Essayer d'extraire le texte brut d'abord
        raw_text = data.get('raw_text', '')
        if raw_text and len(raw_text.strip()) > 10:  # Au moins 10 caractères significatifs
            lines = raw_text.split('\n')
            for i, line in enumerate(lines[:100]):  # Limiter à 100 lignes
                if line.strip():
                    rows.append({
                        'Category': 'Raw_Text',
                        'Type': f'Line_{i+1}',
                        'Content': line.strip()[:500]  # Limiter la longueur
                    })
        
        # Si on a des sections, les ajouter simplement
        sections = data.get('sections', [])
        for section in sections[:10]:
            title = section.get('title', 'Sans titre')
            rows.append({
                'Category': 'Section',
                'Type': 'Title',
                'Content': title
            })
            
            content = section.get('content', [])
            if isinstance(content, list):
                for item in content[:5]:
                    if item and str(item).strip():
                        rows.append({
                            'Category': 'Section',
                            'Type': 'Content',
                            'Content': str(item)[:500]
                        })
        
        # Si toujours vide, analyser pourquoi
        if not rows:
            if raw_text and len(raw_text.strip()) <= 10:
                rows.append({
                    'Category': 'System',
                    'Type': 'Warning',
                    'Content': f'Texte extrait trop court ({len(raw_text)} caractères) - Vérifiez la qualité de l\'image'
                })
            else:
                rows.append({
                    'Category': 'System',
                    'Type': 'Info',
                    'Content': 'Structure de données non reconnue'
                })
        
        print(f"✅ Rows créées: {len(rows)}")
        return pd.DataFrame(rows)
    
    # MÉTHODES DE CONVERSION SPÉCIALISÉES
    def _budget_to_df(self, data):
        """Convertit les données budgétaires en DataFrame"""
        rows = []
        for ligne in data.get('lignes_budgetaires', []):
            rows.append({
                'Description': ligne.get('description', ''),
                'Montant': ligne.get('montant', 0),
                'Type': 'Budget'
            })
        
        if rows and 'total' in data:
            rows.append({
                'Description': 'TOTAL',
                'Montant': data['total'],
                'Type': 'Total'
            })
        
        return pd.DataFrame(rows) if rows else pd.DataFrame({'Message': ['Aucune donnée budgétaire']})
    
    def _lab_to_df(self, data):
        """Convertit les données laboratoire en DataFrame"""
        rows = []
        
        for equip in data.get('equipements', []):
            rows.append({'Element': equip, 'Categorie': 'Equipement'})
        
        for pers in data.get('personnel', []):
            rows.append({'Element': pers, 'Categorie': 'Personnel'})
        
        return pd.DataFrame(rows) if rows else pd.DataFrame({'Message': ['Aucune donnée laboratoire']})
    
    def _rh_laboratoire_to_df(self, data):
        """Convertit les données RH de laboratoire en DataFrame"""
        rows = []
        
        # Métadonnées et statistiques
        if data.get('statistiques'):
            for key, value in data['statistiques'].items():
                rows.append({
                    'CATÉGORIE': 'STATISTIQUES',
                    'TYPE': key,
                    'VALEUR': str(value)
                })
        
        # Personnel par grade
        for item in data.get('personnel_par_grade', []):
            rows.append({
                'CATÉGORIE': 'PERSONNEL',
                'TYPE': 'Grade',
                'VALEUR': item.get('grade', ''),
                'EFFECTIF': item.get('effectif', '')
            })
        
        # Observations
        for i, observation in enumerate(data.get('observations', [])):
            rows.append({
                'CATÉGORIE': 'OBSERVATIONS',
                'TYPE': f'Observation_{i+1}',
                'VALEUR': observation
            })
        
        # Tableaux détectés
        for i, tableau in enumerate(data.get('tableaux', [])):
            rows.append({
                'CATÉGORIE': f'TABLEAU_{i+1}',
                'TYPE': 'Info_Tableau',
                'VALEUR': f"{tableau.get('row_count', 0)} lignes x {tableau.get('column_count', 0)} colonnes"
            })
            
            # En-têtes
            headers = tableau.get('headers', [])
            if headers:
                rows.append({
                    'CATÉGORIE': f'TABLEAU_{i+1}',
                    'TYPE': 'En-têtes',
                    'VALEUR': ' | '.join(headers)
                })
            
            # Données
            for j, row in enumerate(tableau.get('rows', [])):
                row_text = ' | '.join([str(cell) for cell in row])
                rows.append({
                    'CATÉGORIE': f'TABLEAU_{i+1}',
                    'TYPE': f'Ligne_{j+1}',
                    'VALEUR': row_text
                })
        
        return pd.DataFrame(rows) if rows else pd.DataFrame({'Message': ['Aucune donnée RH']})
    
    def _voirie_to_df(self, data):
        """Convertit les données voirie en DataFrame"""
        rows = []
        for troncon in data.get('troncons', []):
            row = {'Description': troncon.get('description', '')}
            if 'longueur' in troncon:
                row['Longueur'] = troncon['longueur']
            if 'largeur' in troncon:
                row['Largeur'] = troncon['largeur']
            if 'type_voirie' in troncon:
                row['Type'] = troncon['type_voirie']
            rows.append(row)
        
        return pd.DataFrame(rows) if rows else pd.DataFrame({'Message': ['Aucune donnée voirie']})
    
    def _formation_to_df(self, data):
        """Convertit les données de formation en DataFrame"""
        rows = []
        
        if data.get('titre'):
            rows.append({'Type': 'Titre', 'Contenu': data['titre']})
        if data.get('date'):
            rows.append({'Type': 'Date', 'Contenu': data['date']})
        
        for instruction in data.get('instructions', []):
            rows.append({'Type': 'Instruction', 'Contenu': instruction})
        
        for exercice in data.get('exercices', []):
            rows.append({'Type': 'Exercice', 'Contenu': exercice.get('titre', '')})
            if exercice.get('description'):
                rows.append({'Type': 'Description', 'Contenu': exercice['description']})
            
            for question in exercice.get('questions', []):
                rows.append({
                    'Type': f"Question {question.get('lettre', '')}",
                    'Contenu': question.get('texte', '')
                })
        
        for question in data.get('questions', []):
            rows.append({
                'Type': f"Question {question.get('lettre', '')}",
                'Contenu': question.get('texte', '')
            })
        
        return pd.DataFrame(rows) if rows else pd.DataFrame({'Message': ['Aucune donnée de formation']})
    
    def _tabular_to_df(self, data):
        """Convertit les données tabulaires en DataFrame amélioré"""
        rows = []
        
        # Métadonnées du document
        if data.get('metadata'):
            for key, values in data['metadata'].items():
                if isinstance(values, list):
                    for value in values[:5]:  # Limiter à 5 valeurs
                        rows.append({
                            'CATÉGORIE': 'METADONNÉES',
                            'TYPE': key,
                            'VALEUR': str(value)
                        })
        
        # Données des tableaux
        for i, table in enumerate(data.get('tables', [])):
            table_num = i + 1
            
            # En-tête du tableau
            rows.append({
                'CATÉGORIE': f'TABLEAU_{table_num}',
                'TYPE': 'EN-TÊTE',
                'VALEUR': f"Tableau {table_num} ({table.get('row_count', 0)} lignes, {table.get('column_count', 0)} colonnes)"
            })
            
            # En-têtes de colonnes
            headers = table.get('headers', [])
            if headers:
                header_row = {f'COLONNE_{j+1}': header for j, header in enumerate(headers)}
                header_row['CATÉGORIE'] = f'TABLEAU_{table_num}'
                header_row['TYPE'] = 'NOM_COLONNES'
                rows.append(header_row)
            
            # Données du tableau
            for j, row_data in enumerate(table.get('rows', [])):
                row_dict = {}
                for k, cell in enumerate(row_data):
                    col_name = headers[k] if k < len(headers) else f'Colonne_{k+1}'
                    row_dict[col_name] = cell
                
                row_dict['CATÉGORIE'] = f'TABLEAU_{table_num}'
                row_dict['TYPE'] = f'LIGNE_{j+1}'
                rows.append(row_dict)
        
        # Données brutes si pas de tableaux détectés
        if not rows and data.get('raw_data'):
            for i, line in enumerate(data.get('raw_data', [])):
                rows.append({
                    'CATÉGORIE': 'DONNÉES_BRUTES',
                    'TYPE': f'Ligne_{i+1}',
                    'CONTENU': line
                })
        
        return pd.DataFrame(rows) if rows else pd.DataFrame({'Message': ['Aucune donnée tabulaire détectée']})
    
    def _legal_to_df(self, data):
        """Convertit les données juridiques en DataFrame"""
        rows = []
        
        for article in data.get('articles', []):
            rows.append({
                'Type': f"Article {article.get('numero', '')}",
                'Titre': article.get('titre', ''),
                'Contenu': ' '.join(article.get('contenu', []))
            })
        
        return pd.DataFrame(rows) if rows else pd.DataFrame({'Message': ['Aucune donnée juridique']})
    
    def _administrative_to_df(self, data):
        """Convertit les données administratives en DataFrame"""
        rows = []
        
        if data.get('expediteur'):
            rows.append({'Type': 'Expéditeur', 'Contenu': data['expediteur']})
        if data.get('destinataire'):
            rows.append({'Type': 'Destinataire', 'Contenu': data['destinataire']})
        if data.get('objet'):
            rows.append({'Type': 'Objet', 'Contenu': data['objet']})
        if data.get('reference'):
            rows.append({'Type': 'Référence', 'Contenu': data['reference']})
        if data.get('date'):
            rows.append({'Type': 'Date', 'Contenu': data['date']})
        
        for i, ligne in enumerate(data.get('contenu', [])):
            rows.append({'Type': f'Contenu_{i+1}', 'Contenu': ligne})
        
        return pd.DataFrame(rows) if rows else pd.DataFrame({'Message': ['Aucune donnée administrative']})
    
    def _generic_to_df(self, data):
        """Convertit les données génériques en DataFrame"""
        # Vérifier d'abord le type spécifique
        if isinstance(data, dict):
            data_type = data.get('type')
            
            # Utiliser les converters spécialisés selon le type
            if data_type == 'budget':
                return self._budget_to_df(data)
            elif data_type == 'laboratoire':
                return self._lab_to_df(data)
            elif data_type == 'rh_laboratoire':
                return self._rh_laboratoire_to_df(data)
            elif data_type == 'voirie':
                return self._voirie_to_df(data)
            elif data_type == 'formation':
                return self._formation_to_df(data)
            elif data_type == 'tabular':
                return self._tabular_to_df(data)
            elif data_type == 'legal':
                return self._legal_to_df(data)
            elif data_type == 'administrative':
                return self._administrative_to_df(data)
        
        # Fallback pour les données génériques
        if isinstance(data, dict) and 'lines' in data:
            lines = data['lines'] or []
            return pd.DataFrame({
                'Ligne': lines,
                'Type': 'Donnee_Generique'
            })
        elif isinstance(data, dict) and 'raw_text' in data:
            raw_text = data['raw_text'] or ''
            lines = [line.strip() for line in raw_text.split('\n') if line.strip()]
            return pd.DataFrame({
                'Ligne': lines,
                'Type': 'Donnee_Generique'
            })
        else:
            return pd.DataFrame({'Message': ['Données non structurées disponibles']})